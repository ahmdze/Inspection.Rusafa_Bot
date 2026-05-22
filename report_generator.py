import os
import re
import sqlite3
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


# =========================================================
# الترتيب الثابت
# =========================================================

DESTINATIONS_ORDER = [
    "الإيعاز الى ادارة المستشفى بما يلي:",
    "الإيعاز الى ادارة القطاع بما يلي:",
    "الإيعاز الى ادارة المركز بما يلي:",
    "الإيعاز الى قسم الامور الادارية والقانونية والمالية بما يلي:",
    "الإيعاز الى شعبة التحقيقات/ قسمنا بما يلي:"
]

AXIS_ORDER = [
    "المعلومات العامة",
    "المحور الفني",
    "المحور الإداري",
    "المحور الهندسي"
]


# =========================================================
# قاعدة البيانات
# =========================================================

def execute_query(query, params=(), fetch=False):
    conn = sqlite3.connect('inspection_db.sqlite')
    cursor = conn.cursor()
    
    try:
        cursor.execute(query, params)
        if fetch:
            return cursor.fetchall()
        conn.commit()
    finally:
        conn.close()


# =========================================================
# أدوات التنسيق
# =========================================================

def force_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    pPr = paragraph._element.get_or_add_pPr()
    
    # إضافة خاصية Right-to-Left بأمان دون تكرار
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.insert(0, bidi)


def set_font(run, size=12, bold=False):
    # التنسيقات الأساسية (المكتبة تتكفل بها بأمان)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'

    rPr = run._element.get_or_add_rPr()
    
    # إضافة علامة RTL للنص إذا لم تكن موجودة
    if rPr.find(qn('w:rtl')) is None:
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)

    # إعداد الخط العربي (Complex Script) دون إنشاء عناصر مكررة
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
        
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(size * 2))

    bCs = rPr.find(qn('w:bCs'))
    if bold:
        if bCs is None:
            bCs = OxmlElement('w:bCs')
            bCs.set(qn('w:val'), '1')
            rPr.append(bCs)
    elif bCs is not None:
        rPr.remove(bCs)


def add_title(doc, text, size=16):
    p = doc.add_paragraph()
    force_rtl(p)
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p


def add_text(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    force_rtl(p)
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold)
    return p


def add_field(doc, label, value):
    p = doc.add_paragraph()
    force_rtl(p)
    
    r1 = p.add_run(f"{label}: ")
    set_font(r1, size=12, bold=True)

    r2 = p.add_run(str(value))
    set_font(r2, size=12)

    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    force_rtl(p)
    
    indent = "    " * level
    run = p.add_run(f"{indent}• {text}")
    set_font(run, size=12)

    return p


# =========================================================
# تنظيف البيانات
# =========================================================

def clean_text(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def format_date(value):
    if not value:
        return ""
    try:
        if isinstance(value, datetime):
            return value.strftime("%d/%m/%Y")
        return datetime.strptime(
            str(value), "%Y-%m-%d"
        ).strftime("%d/%m/%Y")
    except Exception:
        return str(value)


# =========================================================
# تحميل البيانات
# =========================================================

def load_visit_data(visit_id):
    visit_info = execute_query(
        """
        SELECT institution_name, visit_date
        FROM Visits
        WHERE id = ?
        """,
        (visit_id,),
        fetch=True
    )

    if not visit_info:
        return None

    institution_name, visit_date = visit_info[0]

    reports = execute_query(
        """
        SELECT
            axis_name,
            section_name,
            notes,
            rec_destination,
            recommendations,
            user_id
        FROM Reports
        WHERE visit_id = ?
        ORDER BY id ASC
        """,
        (visit_id,),
        fetch=True
    )

    att_result = execute_query(
        """
        SELECT COUNT(*)
        FROM Attachments
        WHERE visit_id = ?
        """,
        (visit_id,),
        fetch=True
    )
    # حماية من أي إرجاع فارغ
    attachments_count = att_result[0][0] if att_result else 0

    members = execute_query(
        """
        SELECT DISTINCT
            user_name
        FROM Visit_Members
        WHERE visit_id = ?
        ORDER BY user_name ASC
        """,
        (visit_id,),
        fetch=True
    )

    grouped_notes = {}
    grouped_recommendations = {}

    for axis, section, note, rec_dest, recs, user_id in (reports or []):
        axis = clean_text(axis)
        section = clean_text(section)
        note = clean_text(note)

        grouped_notes.setdefault(axis, {})
        grouped_notes[axis].setdefault(section, [])
        grouped_notes[axis][section].append(note)

        if rec_dest and recs:
            grouped_recommendations.setdefault(rec_dest, [])
            for line in str(recs).split("\n"):
                line = line.strip()
                if not line:
                    continue
                line = line.lstrip("-").strip()
                grouped_recommendations[rec_dest].append(line)

    return {
        "institution_name": institution_name,
        "visit_date": visit_date,
        "reports": grouped_notes,
        "recommendations": grouped_recommendations,
        "members": [m[0] for m in (members or [])],
        "attachments_count": attachments_count
    }


# =========================================================
# بناء التقرير
# =========================================================

class ReportBuilder:
    def __init__(self, visit_id, data):
        self.visit_id = visit_id
        # تمرير البيانات الجاهزة بدلاً من استدعاء قاعدة البيانات مرة أخرى
        self.data = data  
        self.doc = Document()
        self._setup_document()

    # =====================================================

    def _setup_document(self):
        section = self.doc.sections[0]
        section.top_margin = Pt(50)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(40)
        section.right_margin = Pt(40)

    # =====================================================

    def build_header(self):
        inst = self.data["institution_name"]
        visit_date = format_date(self.data["visit_date"])

        p = self.doc.add_paragraph()
        force_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run("م/ زيارة تفتيشية")
        set_font(run, size=18, bold=True)
        self.doc.add_paragraph()

        intro = (
            "استناداً إلى الخطة السنوية لشعبة تفتيش المؤسسات الصحية الحكومية، "
            f"أجرى فريق من قسم التفتيش زيارة تفتيشية إلى ({inst}) "
            f"بتاريخ ({visit_date}) "
            "وتمت ملاحظة الآتي:"
        )

        p2 = self.doc.add_paragraph()
        force_rtl(p2)
        p2.paragraph_format.line_spacing = 1.5
        run2 = p2.add_run(intro)
        set_font(run2, size=13)
        self.doc.add_paragraph()

    # =====================================================

    def build_notes(self):
        reports = self.data["reports"]
        axis_counter = 1

        for axis in AXIS_ORDER:
            if axis not in reports:
                continue

            sections = reports[axis]
            if not sections:
                continue

            add_title(
                self.doc,
                f"{axis_counter}- {axis}",
                size=16
            )

            section_counter = 1
            for section_name, notes in sections.items():
                if axis == "المعلومات العامة":
                    value = notes[0] if notes else ""
                    add_field(self.doc, section_name, value)
                else:
                    add_title(
                        self.doc,
                        f"{axis_counter}-{section_counter} {section_name}",
                        size=14
                    )
                    for note in notes:
                        add_bullet(self.doc, note, level=1)
                section_counter += 1

            self.doc.add_paragraph()
            axis_counter += 1

    # =====================================================

    def build_recommendations(self):
        recommendations = self.data["recommendations"]
        
        if not recommendations:
            return

        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        force_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        run = p.add_run("الرأي والتوصيات")
        set_font(run, size=18, bold=True)
        self.doc.add_paragraph()

        ordered_destinations = [d for d in DESTINATIONS_ORDER if d in recommendations]
        custom_destinations = [d for d in recommendations.keys() if d not in DESTINATIONS_ORDER]
        ordered_destinations.extend(sorted(custom_destinations))

        for destination in ordered_destinations:
            items = recommendations.get(destination)
            if not items:
                continue

            add_title(self.doc, destination, size=14)

            for idx, item in enumerate(items, start=1):
                p2 = self.doc.add_paragraph()
                force_rtl(p2)
                
                r1 = p2.add_run(f"{idx}- ")
                set_font(r1, size=12, bold=True)
                
                r2 = p2.add_run(item)
                set_font(r2, size=12)

            self.doc.add_paragraph()

    # =====================================================

    def build_statistics(self):
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        force_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        run = p.add_run("إحصائيات الزيارة")
        set_font(run, size=18, bold=True)

        reports_count_res = execute_query(
            """
            SELECT COUNT(*)
            FROM Reports
            WHERE visit_id = ?
            """,
            (self.visit_id,),
            fetch=True
        )
        reports_count = reports_count_res[0][0] if reports_count_res else 0

        self.doc.add_paragraph()

        add_field(self.doc, "عدد الملاحظات", reports_count)
        add_field(self.doc, "عدد المرفقات", self.data["attachments_count"])
        add_field(self.doc, "عدد أعضاء الفريق", len(self.data["members"]))

    # =====================================================

    def build_members(self):
        members = self.data["members"]
        
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        force_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        run = p.add_run("أعضاء الفريق التفتيشي")
        set_font(run, size=18, bold=True)
        self.doc.add_paragraph()

        if not members:
            add_text(self.doc, "لا توجد بيانات أعضاء.")
            return

        for idx, member in enumerate(members, start=1):
            add_bullet(self.doc, f"{idx}- {member}")

    # =====================================================

    def save(self):
        raw_name = str(self.data["institution_name"])
        raw_date = str(self.data["visit_date"])

        # تنظيف الأسماء من أي رموز قد تسبب مشاكل في نظام التشغيل (OS)
        safe_name = re.sub(r'[\\/*?:"<>|]', "-", raw_name)
        safe_date = re.sub(r'[\\/|]', "-", raw_date)

        file_name = f"{safe_name} - {safe_date}.docx"
        self.doc.save(file_name)
        return file_name

    # =====================================================

    def build(self):
        self.build_header()
        self.build_notes()
        self.build_recommendations()
        self.build_statistics()
        self.build_members()
        return self.save()


# =========================================================
# الدالة الرئيسية
# =========================================================

def generate_docx_report(visit_id, bot_token=None):
    data = load_visit_data(visit_id)

    if not data:
        return None, None

    # تمرير البيانات إلى الكلاس لمنع الاستعلام المزدوج
    builder = ReportBuilder(visit_id, data)
    file_name = builder.build()

    return (
        file_name,
        data["institution_name"]
    )